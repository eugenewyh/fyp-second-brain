import logging
from pathlib import Path

from langchain_community.document_loaders import PyMuPDFLoader, TextLoader
from langchain_core.documents import Document

from second_brain.config import SUPPORTED_EXTENSIONS
from second_brain.ingestion.skip import should_skip_ingest_path

logger = logging.getLogger(__name__)


def _extract_docx_text(file_path: Path) -> str:
    from docx import Document as DocxDocument

    document = DocxDocument(str(file_path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def _apply_source_metadata(
    docs: list[Document],
    file_path: Path,
    *,
    ingest_root: Path | None = None,
) -> None:
    resolved = file_path.resolve()
    if ingest_root is not None:
        try:
            source = str(resolved.relative_to(ingest_root.resolve()))
        except ValueError:
            source = file_path.name
    else:
        source = file_path.name

    for doc in docs:
        doc.metadata["source"] = source
        doc.metadata["source_path"] = str(resolved)


def load_file(
    file_path: Path,
    *,
    ingest_root: Path | None = None,
) -> list[Document]:
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        logger.warning("Skipping unsupported file: %s", file_path.name)
        return []

    if suffix == ".pdf":
        loader = PyMuPDFLoader(str(file_path))
        docs = loader.load()
    elif suffix == ".docx":
        try:
            text = _extract_docx_text(file_path)
        except Exception:
            logger.exception("Failed to extract text from %s", file_path.name)
            return []
        if not text:
            logger.warning("No extractable text in %s", file_path.name)
            return []
        docs = [Document(page_content=text, metadata={})]
    else:
        loader = TextLoader(str(file_path), encoding="utf-8")
        docs = loader.load()

    _apply_source_metadata(docs, file_path, ingest_root=ingest_root)

    logger.info("Loaded %d page(s) from %s", len(docs), file_path.name)
    return docs


def load_directory(directory: Path) -> list[Document]:
    if not directory.is_dir():
        raise FileNotFoundError(f"Directory not found: {directory}")

    root = directory.resolve()
    files = sorted(
        p
        for p in root.rglob("*")
        if p.is_file()
        and p.suffix.lower() in SUPPORTED_EXTENSIONS
        and not should_skip_ingest_path(p, root)
    )

    if not files:
        logger.warning("No supported documents found in %s", directory)
        return []

    all_docs: list[Document] = []
    for file_path in files:
        all_docs.extend(load_file(file_path, ingest_root=root))

    logger.info("Loaded %d document(s) from %d file(s)", len(all_docs), len(files))
    return all_docs
