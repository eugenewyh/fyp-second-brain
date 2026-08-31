import sys
import tempfile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from second_brain.ingestion.loaders import load_file
from second_brain.ingestion.pipeline import split_documents


def test_load_text_file():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
        f.write("Hello, Second Brain.")
        path = Path(f.name)

    filename = path.name
    docs = load_file(path)
    path.unlink()

    assert len(docs) == 1
    assert "Second Brain" in docs[0].page_content
    assert docs[0].metadata["source"] == filename


def test_split_documents_adds_metadata():
    from langchain_core.documents import Document

    docs = [Document(page_content="A" * 2000, metadata={"source": "test.txt", "source_path": "/tmp/test.txt"})]
    chunks = split_documents(docs)

    assert len(chunks) >= 2
    assert "chunk_index" in chunks[0].metadata
    assert "ingested_at" in chunks[0].metadata


def test_load_docx_file():
    from docx import Document as DocxDocument

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "notes.docx"
        document = DocxDocument()
        document.add_paragraph("Hello from Word.")
        document.add_paragraph("")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Cell A"
        table.rows[0].cells[1].text = "Cell B"
        document.save(path)

        docs = load_file(path)

        assert len(docs) == 1
        assert "Hello from Word." in docs[0].page_content
        assert "Cell A" in docs[0].page_content
        assert "Cell B" in docs[0].page_content
        assert docs[0].metadata["source"] == "notes.docx"
        assert docs[0].metadata["source_path"] == str(path.resolve())


def test_load_skips_legacy_doc(tmp_path: Path):
    path = tmp_path / "legacy.doc"
    path.write_bytes(b"not a supported word file")
    assert load_file(path) == []


def test_load_directory_recursive_skips_memory(tmp_path: Path):
    from second_brain.ingestion.loaders import load_directory

    topic = tmp_path / "java"
    (topic / "notes").mkdir(parents=True)
    (topic / "memory" / "claims").mkdir(parents=True)
    teachable = topic / "notes" / "a.md"
    teachable.write_text("Nested teachable note.", encoding="utf-8")
    skipped = topic / "memory" / "claims" / "x.md"
    skipped.write_text("Claim file should not ingest.", encoding="utf-8")

    docs = load_directory(tmp_path)

    sources = {d.metadata["source"] for d in docs}
    assert "java/notes/a.md" in sources
    assert not any("memory" in s for s in sources)
    assert "Nested teachable note." in docs[0].page_content


def test_load_pdf_file(tmp_path: Path):
    import fitz

    path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "PyMuPDF extraction works.")
    doc.save(path)
    doc.close()

    docs = load_file(path)

    assert len(docs) >= 1
    assert "PyMuPDF extraction works." in docs[0].page_content

